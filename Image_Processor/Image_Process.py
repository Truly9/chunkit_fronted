import os
import sys
import base64
import io
import re
import concurrent.futures
import hashlib
from typing import List, Dict, Tuple, Optional, Set
from docx import Document
from docx.shared import Inches
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from PIL import Image
import numpy as np
from openai import OpenAI
import json
import time
import pickle
from tqdm import tqdm

current_dir = os.path.dirname(os.path.abspath(__file__)) 
parent_dir = os.path.dirname(current_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

from Utils.Path import (
    PAPER_DOCS_DIR, CAMPUS_DOCS_DIR, FITNESS_DOCS_DIR, PSYCHOLOGY_DOCS_DIR,
    PAPER_INDEX_DIR, CAMPUS_INDEX_DIR, FITNESS_INDEX_DIR, PSYCHOLOGY_INDEX_DIR,
    ALL_PROCESSED_IMAGES_DIR, CAMPUS_IMAGES_DIR, PAPER_IMAGES_DIR, FITNESS_IMAGES_DIR, PSYCHOLOGY_IMAGES_DIR,
    CAMPUS_PROCESSED_EXTRACTED_IMAGES, PSYCHOLOGY_PROCESSED_EXTRACTED_IMAGES,
    CAMPUS_EXTRACTED_IMAGES_JSON, PSYCHOLOGY_EXTRACTED_IMAGES_JSON,
    CAMPUS_IMAGES_PATH, PSYCHOLOGY_IMAGES_PATH
)


class ImageExtractor:
    """从Word文档中提取图片及其上下文的类"""

    def __init__(self, debug_folder: str, output_dir: str = None):
        self.debug_folder = debug_folder
        self.output_dir = output_dir
        
        # 根据输入文件夹自动判断场景
        self.scene = self._detect_scene(debug_folder)
        
        # 如果没有指定输出目录，使用基于场景的默认目录
        if not self.output_dir:
            if self.scene == "campus":
                self.output_dir = str(CAMPUS_IMAGES_PATH)
            elif self.scene == "psychology":
                self.output_dir = str(PSYCHOLOGY_IMAGES_PATH)
            else:
                self.output_dir = os.path.join(os.path.dirname(debug_folder), "extracted_images")
        
        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
        
        print(f"ImageExtractor初始化:")
        print(f"  场景: {self.scene}")
        print(f"  输入目录: {debug_folder}")
        print(f"  输出目录: {self.output_dir}")
        
        # 初始化Qwen3-VL客户端
        self.qwen3_vl_client = OpenAI(
            base_url='https://api-inference.modelscope.cn/v1',
            api_key='ms-03eaf898-af6c-4a07-9857-1afb8337c1b4'
        )

        # 初始化Qwen3.0客户端
        self.qwen3_client = OpenAI(
            base_url='https://api-inference.modelscope.cn/v1',
            api_key='ms-03eaf898-af6c-4a07-9857-1afb8337c1b4'
        )

    def _detect_scene(self, folder_path: str) -> str:
        """根据文件夹路径自动检测场景"""
        folder_path_lower = folder_path.lower()
        if 'campus' in folder_path_lower:
            return "campus"
        elif 'psychology' in folder_path_lower or 'psych' in folder_path_lower:
            return "psychology"
        else:
            # 从路径中提取场景信息
            folder_name = os.path.basename(folder_path)
            if 'campus' in folder_name.lower():
                return "campus"
            elif 'psychology' in folder_name.lower() or 'psych' in folder_name.lower():
                return "psychology"
            else:
                return "unknown"

    def extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """从Word文档中提取图片及其上下文"""
        doc = Document(docx_path)
        images_data = []

        # 获取所有段落文本
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # 遍历文档中的所有段落，查找包含图片的段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                # 查找图片元素
                for drawing in run.element.xpath('.//a:blip'):
                    try:
                        # 获取图片的关系ID
                        embed_id = drawing.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id:
                            # 通过关系ID获取图片数据
                            image_part = doc.part.related_parts[embed_id]
                            image_data = image_part.blob

                            # 获取上下文
                            context_before = ""
                            context_after = ""

                            # 获取图片前后的段落作为上下文
                            if para_idx > 0:
                                context_before = " ".join(
                                    [p.text.strip() for p in doc.paragraphs[max(0, para_idx - 2):para_idx] if
                                     p.text.strip()])[:200]
                            if para_idx < len(doc.paragraphs) - 1:
                                context_after = " ".join([p.text.strip() for p in doc.paragraphs[
                                                                                  para_idx + 1:min(len(doc.paragraphs),
                                                                                                   para_idx + 3)] if
                                                          p.text.strip()])[:200]

                            images_data.append({
                                'image_data': image_data,
                                'context_before': context_before,
                                'context_after': context_after,
                                'source_file': os.path.basename(docx_path),
                                'source_path': docx_path,
                                'paragraph_index': para_idx
                            })
                    except Exception as e:
                        print(f"提取图片时出错: {e}")
                        continue

        return images_data

    def image_to_base64(self, image_data: bytes) -> str:
        """将图片数据转换为base64编码"""
        return base64.b64encode(image_data).decode('utf-8')

    def describe_image_with_qwen3_vl(self, image_data: bytes) -> str:
        """使用Qwen3-VL识别图片内容，无限重试直到成功"""
        attempt = 0
        max_wait_time = 300  # 最大等待5分钟
    
        while True:
            attempt += 1
            try:
                print(f"开始调用Qwen3-VL API... (尝试 {attempt})")
            
                # 将图片转换为base64
                base64_image = self.image_to_base64(image_data)
                image_url = f"data:image/jpeg;base64,{base64_image}"
            
                print(f"图片base64长度: {len(base64_image)}")
                print(f"准备发送请求到模型: Qwen/Qwen3-VL-235B-A22B-Instruct")

                response = self.qwen3_vl_client.chat.completions.create(
                    model='Qwen/Qwen3-VL-235B-A22B-Instruct',
                    messages=[{
                        'role': 'user',
                        'content': [{
                            'type': 'text',
                            'text': '请详细描述这幅图片的内容，包括图片中的文字、图形、布局等所有可见元素。',
                        }, {
                            'type': 'image_url',
                            'image_url': {
                                'url': image_url,
                            },
                        }],
                    }],
                    stream=False,
                )

                result = response.choices[0].message.content
                print(f"Qwen3-VL响应成功，内容长度: {len(result)}")
                return result
            
            except Exception as e:
                error_msg = str(e)
                print(f"Qwen3-VL识别图片时出错 (尝试 {attempt}): {error_msg}")
            
                if "429" in error_msg or "quota" in error_msg.lower():
                    # API限制，等待指数增长的时间
                    wait_time = min(2 ** attempt, max_wait_time)  # 指数退避，最大5分钟
                    print(f"API限流，等待 {wait_time} 秒后重试...")
                    import time
                    time.sleep(wait_time)
                else:
                    # 其他错误，等待较短时间后重试
                    wait_time = min(attempt * 10, 60)  # 线性增长，最大1分钟
                    print(f"API错误，等待 {wait_time} 秒后重试...")
                    import time
                    time.sleep(wait_time)

    def enhance_description_with_qwen3(self, image_description: str, context_before: str, context_after: str) -> str:
        """使用Qwen3.0结合上下文完善图片描述，无限重试直到成功"""
        attempt = 0
        max_wait_time = 300  # 最大等待5分钟
    
        while True:
            attempt += 1
            try:
                print(f"开始调用Qwen3.0 API... (尝试 {attempt})")
            
                prompt = f"""
                请根据以下信息，生成一段完整的图片内容描述：

                图片识别结果：{image_description}

                上文内容：{context_before}

                下文内容：{context_after}

                请结合上下文，分析这张图片的应用场景、主题和作用，并生成一段完整、准确的描述。
                """

                print(f"Qwen3.0提示词长度: {len(prompt)}")
                print(f"准备发送请求到模型: Qwen/Qwen3-235B-A22B-Instruct-2507")

                response = self.qwen3_client.chat.completions.create(
                    model='Qwen/Qwen3-235B-A22B-Instruct-2507',
                    messages=[
                        {
                            'role': 'system',
                            'content': 'You are a helpful assistant that analyzes images in context.'
                        },
                        {
                            'role': 'user',
                            'content': prompt
                        }
                    ],
                    stream=False,
                )

                result = response.choices[0].message.content.strip()
                print(f"Qwen3.0响应成功，内容长度: {len(result)}")
                return result
            
            except Exception as e:
                error_msg = str(e)
                print(f"Qwen3.0完善描述时出错 (尝试 {attempt}): {error_msg}")
            
                if "429" in error_msg or "quota" in error_msg.lower():
                    # API限制，等待指数增长的时间
                    wait_time = min(2 ** attempt, max_wait_time)  # 指数退避，最大5分钟
                    print(f"API限流，等待 {wait_time} 秒后重试...")
                    import time
                    time.sleep(wait_time)
                else:
                    # 其他错误，等待较短时间后重试
                    wait_time = min(attempt * 10, 60)  # 线性增长，最大1分钟
                    print(f"API错误，等待 {wait_time} 秒后重试...")
                    import time
                    time.sleep(wait_time)

    def _generate_image_filename(self, source_file: str, image_index: int, image_data: bytes) -> str:
        """生成统一的图片文件名 - 修复版本"""
        source_name = os.path.splitext(os.path.basename(source_file))[0]
    
        # 使用图片内容生成完整的16位MD5哈希（与图片哈希保持一致）
        image_hash = hashlib.md5(image_data).hexdigest()[:16]
    
        # 统一命名格式：源文件_image_索引_完整16位哈希.jpg
        # 例如：校园邮箱攻略_image_1_5ddf82a0ab52477b.jpg
        filename = f"{source_name}_image_{image_index}_{image_hash}.jpg"
    
        print(f"生成统一文件名: {filename} (哈希: {image_hash})")
        return filename

    def _save_image_file(self, image_data: bytes, filename: str) -> str:
        """保存图片文件到指定目录"""
        image_path = os.path.join(self.output_dir, filename)
        
        try:
            with open(image_path, 'wb') as f:
                f.write(image_data)
            print(f"图片已保存: {filename}")
            return image_path
        except Exception as e:
            print(f"保存图片失败 {filename}: {e}")
            return ""

    def process_all_documents(self, processed_hashes: Set[str] = None) -> List[Dict]:
        """处理debug文件夹中的所有Word文档（包括子文件夹），支持增量处理"""
        if processed_hashes is None:
            processed_hashes = set()
            
        all_processed_images = []
        json_file_path = self._get_scene_json_path()

        # 递归获取所有docx文件
        docx_files = []
        for root, dirs, files in os.walk(self.debug_folder):
            for file in files:
                # 跳过临时文件
                if file.startswith('~$') or file.startswith('.') or file in ['Thumbs.db', '.DS_Store']:
                    continue
                if file.endswith('.docx'):
                    docx_files.append(os.path.join(root, file))
    
        print(f"找到 {len(docx_files)} 个Word文档（包括子文件夹）")
    
        for docx_path in docx_files:
            print(f"\n处理文档: {docx_path}")
        
            try:
                # 提取图片
                images_data = self.extract_images_from_docx(docx_path)
                print(f"从 {os.path.basename(docx_path)} 中提取到 {len(images_data)} 张图片")
    
                for i, img_data in enumerate(images_data):
                    # 计算图片哈希
                    image_hash = hashlib.md5(img_data['image_data']).hexdigest()[:16]
                    
                    # 检查图片是否已处理过
                    if image_hash in processed_hashes:
                        print(f"  跳过已处理的图片 {i+1}/{len(images_data)} (哈希: {image_hash})")
                        continue
                    
                    print(f"  处理图片 {i + 1}/{len(images_data)}...")
    
                    # 生成统一的文件名 - 使用完整的16位哈希
                    source_file = img_data['source_file']
                    filename = self._generate_image_filename(source_file, i + 1, img_data['image_data'])
                    
                    # 保存图片文件
                    image_path = self._save_image_file(img_data['image_data'], filename)
                    
                    if not image_path:
                        continue  # 如果保存失败，跳过这张图片
    
                    # 使用Qwen3-VL识别图片（无限重试）
                    try:
                        image_description = self.describe_image_with_qwen3_vl(img_data['image_data'])
                        print(f"    图片识别完成")
                    except Exception as e:
                        print(f"    ❌ 图片识别失败: {e}")
                        image_description = f"图片识别失败: {str(e)}"
    
                    # 使用Qwen3.0完善描述（无限重试）
                    try:
                        enhanced_description = self.enhance_description_with_qwen3(
                            image_description,
                            img_data['context_before'],
                            img_data['context_after']
                        )
                        print(f"    描述完善完成")
                    except Exception as e:
                        print(f"    ❌ 描述完善失败: {e}")
                        enhanced_description = f"{image_description} [描述完善失败: {str(e)}]"
    
                    # 构建处理后的图片数据
                    processed_img = {
                        'context_before': img_data['context_before'],
                        'context_after': img_data['context_after'],
                        'source_file': source_file,
                        'source_path': img_data['source_path'],
                        'image_path': image_path,
                        'image_filename': filename,
                        'saved_path': image_path,
                        'processed_path': image_path,
                        'original_description': image_description,
                        'enhanced_description': enhanced_description,
                        'image_hash': image_hash,  # 使用从文件名提取的哈希
                        'scene': self.scene,
                        'paragraph_index': img_data.get('paragraph_index', 0)
                    }
                    
                    # 添加到结果列表
                    all_processed_images.append(processed_img)
                    
                    # 立即写入JSON文件
                    self._append_to_json(processed_img, json_file_path)
                    
                    # 添加到已处理哈希集合
                    processed_hashes.add(image_hash)
                    
                    # 增加延迟避免API限制
                    delay = 5  # 统一5秒延迟
                    print(f"    等待 {delay} 秒避免API限制...")
                    time.sleep(delay)

            except Exception as e:
                print(f"处理文档 {docx_path} 时出错: {e}")
                continue
        
        print(f"\n图片处理完成: 共处理 {len(all_processed_images)} 张图片")
        print(f"   输出目录: {self.output_dir}")
        return all_processed_images

    def _get_scene_json_path(self) -> str:
        """获取场景对应的JSON文件路径"""
        if self.scene == "campus":
            return str(CAMPUS_EXTRACTED_IMAGES_JSON)
        elif self.scene == "psychology":
            return str(PSYCHOLOGY_EXTRACTED_IMAGES_JSON)
        else:
            return os.path.join(self.output_dir, f"{self.scene}_extracted_images.json")

    def _append_to_json(self, img_data: Dict, json_path: str):
        """将图片数据追加到JSON文件"""
        # 确保文件存在
        if not os.path.exists(json_path):
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump([], f, ensure_ascii=False, indent=2)
        
        # 读取现有数据
        existing_data = []
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                existing_data = json.load(f)
        except Exception as e:
            print(f"读取JSON文件时出错: {e}")
            existing_data = []
        
        # 添加新数据
        existing_data.append(img_data)
        
        # 写回文件
        try:
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(existing_data, f, ensure_ascii=False, indent=2)
            print(f"图片数据已追加到JSON文件: {json_path}")
        except Exception as e:
            print(f"写入JSON文件失败: {e}")

    def _verify_filename_consistency(self, processed_images: List[Dict]):
        """验证文件名一致性"""
        print("\n=== 验证文件名一致性 ===")
    
        consistent_count = 0
        inconsistent_count = 0
    
        for img_data in processed_images:
            filename = img_data.get('image_filename', '')
            stored_hash = img_data.get('image_hash', '')
        
            # 从文件名中提取哈希
            if '_image_' in filename:
                # 匹配_image_后面的哈希值（假设哈希是16位十六进制数）
                match = re.search(r'_image_.*?([a-f0-9]{16})\.', filename)
                if match:
                    hash_from_filename = match.group(1)
                    # 检查哈希是否一致
                    if hash_from_filename == stored_hash and len(stored_hash) == 16:
                        consistent_count += 1
                    else:
                        inconsistent_count += 1
                        print(f"❌ 不一致: 文件名哈希={hash_from_filename}, 存储哈希={stored_hash}")
    
        print(f"一致性检查结果:")
        print(f"一致的文件: {consistent_count}")
        print(f"不一致的文件: {inconsistent_count}")
    
        if inconsistent_count == 0:
            print("所有文件名和哈希完全一致！")
        else:
            print("存在不一致的文件，需要修复")

    def save_images_to_word(self, processed_images: List[Dict], output_path: str):
        """将处理后的图片和描述保存到Word文档"""
        doc = Document()
        doc.add_heading('提取的图片及其描述', 0)

        for i, img_data in enumerate(processed_images):
            # 添加图片标题
            doc.add_heading(f'图片 {i + 1} - 来源: {img_data["source_file"]}', level=1)

            # 添加图片
            try:
                # 将图片数据保存为临时文件
                temp_image_path = f"temp_image_{i}.png"
                with open(temp_image_path, 'wb') as f:
                    f.write(img_data['image_data'])

                # 添加图片到文档
                doc.add_picture(temp_image_path, width=Inches(4))

                # 删除临时文件
                os.remove(temp_image_path)
            except Exception as e:
                doc.add_paragraph(f"图片加载失败: {e}")

            # 添加上下文
            doc.add_heading('上下文', level=2)
            doc.add_paragraph(f"上文: {img_data['context_before']}")
            doc.add_paragraph(f"下文: {img_data['context_after']}")

            # 添加描述
            doc.add_heading('图片描述', level=2)
            doc.add_paragraph(img_data['enhanced_description'])

            # 添加分隔线
            doc.add_paragraph('\n' + '=' * 50 + '\n')

        doc.save(output_path)
        print(f"图片和描述已保存到: {output_path}")


class UnifiedImageManager:
    """统一的图片管理器，确保MultiRAG和ImageExtractor使用相同的路径和命名规则"""
    
    def __init__(self, scene: str):
        self.scene = scene
        
        # 根据场景设置路径
        if scene == "campus":
            self.docs_dir = str(CAMPUS_DOCS_DIR)
            self.output_dir = str(CAMPUS_IMAGES_PATH)
        else:
            self.docs_dir = str(PSYCHOLOGY_DOCS_DIR) 
            self.output_dir = str(PSYCHOLOGY_IMAGES_PATH)
            
        self.extractor = ImageExtractor(self.docs_dir, self.output_dir)
        self._ensure_directories()
    
    def _ensure_directories(self):
        """确保目录存在"""
        os.makedirs(self.output_dir, exist_ok=True)
    
    def process_images(self, incremental: bool = True, force_reprocess: bool = False) -> List[Dict]:
        """处理图片并返回格式化数据，支持增量处理"""
        # 加载已处理的图片哈希
        processed_hashes = self._load_processed_hashes()
        
        # 如果需要强制重新处理，清空已处理哈希
        if force_reprocess:
            print(f"强制重新处理场景: {self.scene}")
            processed_hashes = set()
            
            # 删除JSON文件以重新开始
            json_path = self._get_scene_json_path()
            if os.path.exists(json_path):
                os.remove(json_path)
                print(f"已删除JSON文件: {json_path}")
        
        # 使用ImageExtractor处理图片
        new_images = self.extractor.process_all_documents(processed_hashes)
        return new_images
    
    def _load_processed_hashes(self) -> Set[str]:
        """从JSON文件加载已处理的图片哈希"""
        json_path = self._get_scene_json_path()
        
        if not os.path.exists(json_path):
            return set()
        
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                hashes = set(img['image_hash'] for img in data)
                print(f"从JSON文件加载了 {len(hashes)} 个已处理的图片哈希")
                return hashes
        except Exception as e:
            print(f"加载已处理哈希时出错: {e}")
            return set()
    
    def _get_scene_json_path(self) -> str:
        """获取场景对应的JSON文件路径"""
        if self.scene == "campus":
            return str(CAMPUS_EXTRACTED_IMAGES_JSON)
        elif self.scene == "psychology":
            return str(PSYCHOLOGY_EXTRACTED_IMAGES_JSON)
        else:
            return os.path.join(self.output_dir, f"{self.scene}_extracted_images.json")
    
    def reset_scene(self):
        """重置场景处理进度（删除JSON文件）"""
        json_path = self._get_scene_json_path()
        if os.path.exists(json_path):
            os.remove(json_path)
            print(f"已重置场景 {self.scene} 的处理进度")
        else:
            print(f"场景 {self.scene} 没有找到可重置的进度文件")


def main():
    """主函数 - 处理多个场景"""
    # 处理 campus 场景
    print("=" * 50)
    print("开始处理 CAMPUS 场景的Word文档中的图片...")
    
    campus_manager = UnifiedImageManager("campus")
    processed_images_campus = campus_manager.process_images(incremental=True)
    
    if processed_images_campus:
        print(f"\nCAMPUS 场景总共处理了 {len(processed_images_campus)} 张图片")
    else:
        print("CAMPUS 场景没有找到任何新图片")
    
    #处理 psychology 场景（强制重新处理）
    print("\n" + "=" * 50)
    print("开始处理 PSYCHOLOGY 场景的Word文档中的图片...")
    
    # 重置psychology场景的处理进度
    psychology_manager = UnifiedImageManager("psychology")
    psychology_manager.reset_scene()
    
    # 强制重新处理所有图片
    processed_images_psychology = psychology_manager.process_images(force_reprocess=True)
    
    if processed_images_psychology:
        print(f"\nPSYCHOLOGY 场景总共处理了 {len(processed_images_psychology)} 张图片")
    else:
        print("PSYCHOLOGY 场景没有找到任何图片")
    
    print("\n" + "=" * 50)
    print("所有场景的图片提取和处理完成！")

if __name__ == "__main__":
    main()