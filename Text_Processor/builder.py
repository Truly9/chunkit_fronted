#----文档 处理和向量化的主要逻辑----#

import tiktoken
import numpy as np
import os
import sys
import torch
import PyPDF2
import hashlib
import json
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document

from faiss_store_y import FAISSVectorStore
from sentence_transformers import SentenceTransformer
from textsplitters import RecursiveCharacterTextSplitter

current_dir = os.path.dirname(os.path.abspath(__file__)) 
parent_dir = os.path.dirname(current_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)


from Utils.Path import (
    PAPER_DOCS_DIR, CAMPUS_DOCS_DIR, FITNESS_DOCS_DIR, PSYCHOLOGY_DOCS_DIR,
    PAPER_INDEX_DIR, CAMPUS_INDEX_DIR, FITNESS_INDEX_DIR, PSYCHOLOGY_INDEX_DIR
)

# 配置文件路径

#PROCESS_STATUS_FILE 记录的是每个知识库中每个文件的文本处理状态（包括文件哈希、处理时间等）
## 和  <最后一次文本处理时>  的图像处理状态 ,因此不一定是最新的图像处理状态，最新的在 以下各个图像处理状态文件
PROCESS_STATUS_FILE = os.path.join(current_dir, "process_status.json")
BACKUP_DIR = os.path.join(current_dir, "index_backups")

# 图像处理状态文件路径
Base_IMAGE_FILE = os.path.join(parent_dir, "All_Processed_Images")
PSYCHOLOGY_IMAGES_FILE = os.path.join(Base_IMAGE_FILE, "psychology", "psychology_extracted_images.json")
CAMPUS_IMAGES_FILE = os.path.join(Base_IMAGE_FILE, "campus", "campus_extracted_images.json")
FITNESS_IMAGES_FILE = os.path.join(Base_IMAGE_FILE, "fitness", "fitness_extracted_images.json")
PAPER_IMAGES_FILE = os.path.join(Base_IMAGE_FILE, "paper", "paper_extracted_images.json")

def create_backup(index_path, collection_name):
    """创建索引备份"""
    os.makedirs(BACKUP_DIR, exist_ok=True)
    if os.path.exists(index_path):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(BACKUP_DIR, f"{collection_name}_backup_{timestamp}")
        shutil.copytree(index_path, backup_path)
        print(f"已创建备份: {backup_path}")
        return backup_path
    return None

def load_process_status():
    """加载处理状态记录"""
    if os.path.exists(PROCESS_STATUS_FILE):
        try:
            with open(PROCESS_STATUS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"读取处理状态文件失败: {str(e)}")
            return {}
    return {}

def save_process_status(status):
    """保存处理状态记录"""
    try:
        with open(PROCESS_STATUS_FILE, 'w', encoding='utf-8') as f:
            json.dump(status, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"保存处理状态文件失败: {str(e)}")

def load_images_status(images_file_path):
    """加载图像处理状态"""
    if os.path.exists(images_file_path):
        try:
            with open(images_file_path, 'r', encoding='utf-8') as f:
                images_data = json.load(f)
                # 转换为以source_file为键的字典，方便查找
                images_status = {}
                for item in images_data:
                    if isinstance(item, dict) and 'source_file' in item:
                        images_status[item['source_file']] = True  #记录的图片文件路径存在，就表示已处理
                return images_status
        except Exception as e:
            print(f"读取图像处理状态文件失败 {images_file_path}: {str(e)}")
            return {}
    return {}

def get_file_hash(filepath):
    """计算文件的MD5哈希值"""
    try:
        with open(filepath, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except Exception as e:
        print(f"计算文件哈希失败 {filepath}: {str(e)}")
        return None

def read_file(filename: str) -> str:
    """读取文件内容，支持多种编码格式"""
    if not os.path.exists(filename):
        raise FileNotFoundError(f"文件不存在: {filename}")
    
    file_size = os.path.getsize(filename)
    if file_size == 0:
        raise ValueError(f"文件为空: {filename}")
    if file_size > 100 * 1024 * 1024:  # 100MB限制
        raise ValueError(f"文件过大 ({file_size} bytes): {filename}")

    try:
        if filename.endswith('.txt'):
            encodings = ['utf-8', 'gbk', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(filename, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():  # 确保内容非空
                            return content
                except UnicodeDecodeError:
                    continue
            raise Exception(f"无法解码txt文件: {filename}")

        elif filename.endswith(('.md', '.markdown')):
            with open(filename, 'r', encoding='utf-8') as file:
                return file.read()

        elif filename.endswith('.pdf'):
            text = ""
            with open(filename, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                if len(reader.pages) == 0:
                    raise ValueError("PDF文件没有页面")
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                raise ValueError("PDF文件没有提取到文本内容")
            return text

        elif filename.endswith('.docx'):
            doc = Document(filename)
            if len(doc.paragraphs) == 0:
                raise ValueError("Word文档没有内容")
            
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            if not text.strip():
                raise ValueError("Word文档没有文本内容")
            return text

        else:
            raise ValueError(f"不支持的文件格式: {filename}")

    except Exception as e:
        raise Exception(f"读取文件 {filename} 失败: {str(e)}")

def token_length_function(text: str) -> int:
    """计算文本的token长度"""
    try:
        encoding = tiktoken.get_encoding('cl100k_base')
        return len(encoding.encode(text))
    except Exception as e:
        raise Exception(f"计算token长度失败: {str(e)}")

def split_document(filename: str) -> list[str]:
    """分割文档为多个chunk"""
    try:
        content = read_file(filename)
        if not content.strip():
            raise ValueError("文件内容为空，无法进行分割")

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=300,
            chunk_overlap=50,
            length_function=token_length_function,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " "],
        )
        
        chunks = text_splitter.split_text(content)

        if not chunks:
            raise ValueError("文档分割后未得到任何片段")

        print(f"文档 {os.path.basename(filename)} 分割完成，共得到 {len(chunks)} 个片段")
        return chunks

    except Exception as e:
        raise Exception(f"分割文档 {filename} 失败: {str(e)}")

class BaseAgentKnowledgeBase:
    """所有智能体知识库的抽象基类"""

    def __init__(self, index_path: str, collection_name: str, dimension: int = 1024, images_status_file: str = None):
        self.index_path = index_path
        self.collection_name = collection_name
        self.dimension = dimension
        self.images_status_file = images_status_file
        self.process_status = load_process_status()
        
        # 加载图像处理状态
        self.images_status = {}
        if self.images_status_file and os.path.exists(self.images_status_file):
            self.images_status = load_images_status(self.images_status_file)
            print(f"加载图像处理状态，已处理 {len(self.images_status)} 个文档的图像")
        
        # 确保索引目录存在
        os.makedirs(index_path, exist_ok=True)
        
        # 创建备份
        self.backup_path = create_backup(index_path, collection_name)
        
        # 安全初始化向量存储
        self.vector_store = self._init_faiss_store_safe()
        self.model = self._load_embedding_model()

    def _init_faiss_store_safe(self):
        """安全初始化FAISS向量存储"""
        try:
            # 检查是否已有索引文件
            has_existing_index = False
            if os.path.exists(self.index_path):
                index_files = [f for f in os.listdir(self.index_path) if f.endswith(('.faiss', '.pkl', '.json'))]
                has_existing_index = len(index_files) > 0
            
            if has_existing_index:
                print(f"检测到现有FAISS索引，将加载现有索引: {self.collection_name}")
                reset = False
            else:
                print(f"未找到现有FAISS索引，将创建新索引: {self.collection_name}")
                reset = False  # 即使没有索引也不重置，让FAISSStore自己处理
            
            vector_store = FAISSVectorStore(
                index_path=self.index_path,
                collection_name=self.collection_name,
                dimension=self.dimension,
                reset=False  # 强制设置为False，确保不重置
            )
            
            current_count = vector_store.count()
            print(f"FAISS向量存储初始化成功，集合名: {self.collection_name}，当前文档数量: {current_count}")
            return vector_store
            
        except Exception as e:
            print(f"初始化FAISS向量存储失败: {str(e)}")
            # 尝试使用reset=False重新初始化
            try:
                print("尝试使用安全模式重新初始化...")
                vector_store = FAISSVectorStore(
                    index_path=self.index_path,
                    collection_name=self.collection_name,
                    dimension=self.dimension,
                    reset=False
                )
                print(f"安全模式初始化成功，文档数量: {vector_store.count()}")
                return vector_store
            except Exception as e2:
                raise Exception(f"FAISS向量存储初始化完全失败: {str(e)}, 安全模式也失败: {str(e2)}")

    def _load_embedding_model(self):
        """加载嵌入模型"""
        device = "cuda" if torch.cuda.is_available() else "cpu"
        print(f"使用设备: {device}")

        try:
            # 尝试多种可能的模型路径
            possible_paths = [
                "Qwen3-Embedding-0___6B",
                "..Qwen3-Embedding-0___6B"
            ]
            
            model_path = None
            for path in possible_paths:
                abs_path = os.path.abspath(path)
                if os.path.exists(abs_path):
                    model_path = abs_path
                    break
            
            if not model_path:
                raise FileNotFoundError("未找到模型目录，请检查模型路径")
            
            print(f"使用模型路径: {model_path}")
            
            model = SentenceTransformer(
                model_path,
                device=device,
                trust_remote_code=True
            )
            
            # 测试模型
            test_embedding = model.encode(["测试文本"])
            print(f"嵌入模型加载成功，测试向量维度: {test_embedding.shape}")
            
            return model
            
        except Exception as e:
            print(f"加载嵌入模型失败: {str(e)}")


    def _get_collection_status_key(self):
        """获取当前集合的状态键"""
        return f"{self.collection_name}_status"

    def _is_file_processed(self, file_path: str) -> bool:
        """检查文件是否已经处理过（文本处理）"""
        status_key = self._get_collection_status_key()
        if status_key not in self.process_status:
            return False
            
        file_hash = get_file_hash(file_path)
        if not file_hash:
            return False
            
        file_status = self.process_status[status_key].get(file_path, {})
        return file_status.get('hash') == file_hash and file_status.get('processed', False)#如果 'processed' 键不存在，则返回默认值 False

    def _is_file_images_processed(self, file_path: str) -> bool:
        """检查文件的图像是否已经处理过"""
        if not self.images_status:
            return False
            
        filename = os.path.basename(file_path)
        return filename in self.images_status

    def _mark_file_processed(self, file_path: str):
        """标记文件为已处理"""
        status_key = self._get_collection_status_key()
        if status_key not in self.process_status:
            self.process_status[status_key] = {}
            
        file_hash = get_file_hash(file_path)
        if file_hash:
            self.process_status[status_key][file_path] = {
                'hash': file_hash,
                'processed': True,
                'last_processed': datetime.now().isoformat(),
                'images_processed': self._is_file_images_processed(file_path)
            }
            save_process_status(self.process_status)

    def process_folder(self, folder_name: str, force_reprocess: bool = False, skip_images_processed: bool = False):
        """处理指定文件夹中的文件，支持增量处理和图像处理状态检测"""
        if not os.path.isdir(folder_name):
            raise NotADirectoryError(f"文件夹 {folder_name} 不存在")

        print(f"开始处理文件夹: {folder_name}")
        print(f"当前向量库文档数量: {self.vector_store.count()}")

        # 获取支持的文件
        supported_extensions = {'.txt', '.md', '.markdown', '.pdf', '.docx'}
        all_files = []
        for root, dirs, files in os.walk(folder_name):
            for file in files:
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in supported_extensions:
                    full_path = os.path.join(root, file)
                    all_files.append(full_path)

        if not all_files:
            print(f"文件夹 {folder_name} 中没有支持的文件")
            return

        print(f"发现 {len(all_files)} 个支持的文件")

        # 过滤已处理的文件
        if not force_reprocess:
            unprocessed_files = []
            for file_path in all_files:
                filename = os.path.basename(file_path)
                
                # 检查文本是否已处理
                text_processed = self._is_file_processed(file_path)
                
                # 检查图像是否已处理
                images_processed = self._is_file_images_processed(file_path)
                
                if text_processed:
                    if images_processed and skip_images_processed:
                        print(f"跳过已完全处理文件(文本+图像): {filename}")
                        continue
                    elif not force_reprocess:
                        print(f"跳过已处理文件(文本): {filename}")
                        continue
                
                unprocessed_files.append(file_path)
            
            all_files = unprocessed_files
            print(f"需要处理的文件数量: {len(all_files)}")

        if not all_files:
            print("所有文件都已处理完成，无需重复处理")
            return

        total_chunks = 0
        successful_files = 0
        
        print(f"\n开始处理 {len(all_files)} 个文件...")
        
        for file_idx, file_path in enumerate(all_files, 1):
            try:
                filename = os.path.basename(file_path)
                print(f"\n[{file_idx}/{len(all_files)}] 处理文件: {filename}")
                
                # 检查图像处理状态
                images_processed = self._is_file_images_processed(file_path)
                if images_processed:
                    print(f"此文件的图像已被处理过")
                
                # 处理文档
                chunks = split_document(file_path)
                total_chunks += len(chunks)

                # 批量编码
                print(f"  正在编码 {len(chunks)} 个文本片段...")
                embeddings = self.model.encode(chunks, batch_size=32, show_progress_bar=False)
                
                documents = []
                embeddings_list = []
                ids = []

                for chunk_idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    unique_id = f"{hashlib.md5(file_path.encode()).hexdigest()[:8]}_chunk_{chunk_idx}"
                    documents.append(chunk)
                    embeddings_list.append(embedding.tolist())
                    ids.append(unique_id)

                # 添加到向量库
                print(f"  正在添加到向量库...")
                self.vector_store.add(
                    documents=documents,
                    embeddings=embeddings_list,
                    ids=ids
                )
                
                # 标记为已处理
                self._mark_file_processed(file_path)
                successful_files += 1
                print(f"文件处理完成，新增 {len(chunks)} 个片段")

            except Exception as e:
                print(f"处理文件失败: {str(e)}")
                import traceback
                traceback.print_exc()
                continue

        print(f"\n处理完成统计:")
        print(f"成功处理: {successful_files}/{len(all_files)} 个文件")
        print(f"生成片段: {total_chunks} 个")
        print(f"当前向量库文档总数: {self.vector_store.count()}")

    def get_processing_status(self):
        """获取处理状态统计"""
        status_key = self._get_collection_status_key()
        if status_key not in self.process_status:
            return {
                "total_processed": 0, 
                "total_files": 0,
                "images_processed": 0,
                "completion_rate": "0%",
                "files": {}
            }
        
        files_status = self.process_status[status_key]
        processed_count = 0
        images_processed_count = 0
        
        for file_path, status in files_status.items():
            if status.get('processed', False):
                processed_count += 1
                if status.get('images_processed', False):
                    images_processed_count += 1
        
        completion_rate = "0%"
        if files_status:
            completion_rate = f"{(processed_count/len(files_status)*100):.1f}%"
        
        return {
            "total_processed": processed_count,
            "total_files": len(files_status),
            "images_processed": images_processed_count,
            "completion_rate": completion_rate,
            "files": files_status
        }

    def restore_from_backup(self):
        """从备份恢复索引"""
        if self.backup_path and os.path.exists(self.backup_path):
            print(f"正在从备份恢复: {self.backup_path}")
            # 清空当前索引目录
            if os.path.exists(self.index_path):
                shutil.rmtree(self.index_path)
            # 从备份恢复
            shutil.copytree(self.backup_path, self.index_path)
            print("恢复完成")
            # 重新初始化向量存储
            self.vector_store = self._init_faiss_store_safe()
        else:
            print("没有可用的备份")

# 子类定义
class PsychologyAssistant(BaseAgentKnowledgeBase):
    def __init__(self):
        super().__init__(
            index_path=str(PSYCHOLOGY_INDEX_DIR),
            collection_name="psychology_docs",
            dimension=1024,
            images_status_file=PSYCHOLOGY_IMAGES_FILE
        )

class CampusQnA(BaseAgentKnowledgeBase):
    def __init__(self):
        super().__init__(
            index_path=str(CAMPUS_INDEX_DIR),
            collection_name="campus_docs",
            dimension=1024,
            images_status_file=CAMPUS_IMAGES_FILE
        )

class FitnessDietAssistant(BaseAgentKnowledgeBase):
    def __init__(self):
        super().__init__(
            index_path=str(FITNESS_INDEX_DIR),
            collection_name="fitness_docs",
            dimension=1024,
            images_status_file=FITNESS_IMAGES_FILE
        )

class PaperAssistant(BaseAgentKnowledgeBase):
    def __init__(self):
        super().__init__(
            index_path=str(PAPER_INDEX_DIR),
            collection_name="paper_docs",
            dimension=1024,
            images_status_file=PAPER_IMAGES_FILE
        )


if __name__ == '__main__':
    try:
        print("=== 知识库重建工具 ===")
        print("检测到FAISS索引丢失，正在重新构建...")
        
        # 显示当前状态
        psychology_kb = PsychologyAssistant()
        campus_kb = CampusQnA()
        fitness_kb = FitnessDietAssistant()
        paper_kb = PaperAssistant()
        
        knowledge_bases = [
            ("心理助手", psychology_kb, PSYCHOLOGY_DOCS_DIR),
            ("校园问答", campus_kb, CAMPUS_DOCS_DIR),
            ("健身饮食", fitness_kb, FITNESS_DOCS_DIR),
            ("论文助手", paper_kb, PAPER_DOCS_DIR)
        ]
        
        print("\n当前索引状态:")
        for name, kb, _ in knowledge_bases:
            status = kb.get_processing_status()
            count = kb.vector_store.count()
            print(f"  {name}: {count} 个文档, 完成率: {status['completion_rate']}")
        
        print("\n" + "="*50)
        
        # 询问重建选项
        print("\n重建选项:")
        print("1. 智能重建（推荐）：只处理未处理的文件")
        print("2. 完全重建：重新处理所有文件")
        print("3. 跳过图像已处理文件：只处理图像未处理的文件")
        
        user_choice = input("请选择重建模式 (1/2/3, 默认1): ").strip()
        
        if user_choice == '2':
            force_reprocess = True
            skip_images_processed = False
            print("完全重建模式：重新处理所有文件")
        elif user_choice == '3':
            force_reprocess = False
            skip_images_processed = True
            print("跳过图像已处理文件模式")
        else:
            force_reprocess = False
            skip_images_processed = False
            print("智能重建模式：只处理未处理的文件")
        
        print("\n" + "="*50)
        print("开始重建FAISS索引...")
        
        # 重建各知识库
        for name, kb, docs_dir in knowledge_bases:
            print(f"\n--- 重建{name}知识库 ---")
            try:
                kb.process_folder(
                    folder_name=str(docs_dir), 
                    force_reprocess=force_reprocess,
                    skip_images_processed=skip_images_processed
                )
            except Exception as e:
                print(f"重建{name}知识库时出错: {str(e)}")
                import traceback
                traceback.print_exc()
                continue
        
        print("\n=== FAISS索引重建完成！ ===")
        
        # 显示最终状态
        print("\n最终索引状态:")
        for name, kb, _ in knowledge_bases:
            status = kb.get_processing_status()
            count = kb.vector_store.count()
            print(f"  {name}: {count} 个文档, 完成率: {status['completion_rate']}")
        
        print(f"\n备份文件保存在: {BACKUP_DIR}")
        
    except Exception as e:
        print(f"程序执行失败: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)